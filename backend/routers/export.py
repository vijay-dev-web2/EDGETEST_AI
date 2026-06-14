from __future__ import annotations

import io
import json
import pathlib
import uuid

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from github import Github, GithubException
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from auth import get_current_user
from database import get_db
from models import Session as SessionModel, TestRun, User

router = APIRouter()

_TEMPLATES_DIR = pathlib.Path(__file__).parent.parent / "templates"
_jinja_env = Environment(loader=FileSystemLoader(str(_TEMPLATES_DIR)), autoescape=False)

_LANG_DEFAULTS: dict[str, dict[str, str]] = {
    "python":     {"dependency_file": "requirements.txt", "test_glob": "test_*.py"},
    "javascript": {"dependency_file": "package.json",     "test_glob": "*.test.js"},
    "typescript": {"dependency_file": "package.json",     "test_glob": "*.test.ts"},
}


class WorkflowExportRequest(BaseModel):
    language: str
    test_file_paths: list[str]
    dependency_file: str
    session_id: str | None = None
    coverage_threshold: int | None = None


class PushToGithubRequest(BaseModel):
    session_id: uuid.UUID
    repo_full_name: str


def _render_workflow(
    language: str,
    test_paths: list[str],
    dependency_file: str,
    coverage_threshold: int | None = None,
) -> str:
    if language == "java":
        template_name = "github_actions_java.yml.j2"
    elif language == "csharp":
        template_name = "github_actions_csharp.yml.j2"
    else:
        template_name = "github_actions.yml.j2"
    template = _jinja_env.get_template(template_name)
    return template.render(
        language=language,
        test_paths=test_paths,
        dependency_file=dependency_file,
        coverage_threshold=coverage_threshold,
    )


@router.post("/workflow")
async def export_workflow(payload: WorkflowExportRequest) -> Response:
    yaml_content = _render_workflow(
        payload.language,
        payload.test_file_paths,
        payload.dependency_file,
        payload.coverage_threshold,
    )
    return Response(
        content=yaml_content,
        media_type="application/x-yaml",
        headers={"Content-Disposition": "attachment; filename=github-actions-tests.yml"},
    )


@router.post("/push-to-github")
async def push_to_github(
    payload: PushToGithubRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    session = await db.get(SessionModel, payload.session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    language = session.language
    defaults = _LANG_DEFAULTS.get(language, _LANG_DEFAULTS["python"])
    test_dir = pathlib.Path("/tmp/edgetest") / str(payload.session_id)
    test_paths = sorted(f.name for f in test_dir.iterdir() if f.is_file()) if test_dir.is_dir() else [defaults["test_glob"]]

    yaml_content = _render_workflow(language, test_paths, defaults["dependency_file"])
    workflow_path = ".github/workflows/test-generator.yml"
    try:
        g = Github(current_user.access_token)
        repo = g.get_repo(payload.repo_full_name)
        sha: str | None = None
        try:
            existing = repo.get_contents(workflow_path)
            sha = existing.sha  # type: ignore[union-attr]
        except GithubException as exc:
            if exc.status != 404:
                raise
        if sha is not None:
            repo.update_file(workflow_path, "chore: update EdgeTest AI generated workflow", yaml_content, sha)
        else:
            repo.create_file(workflow_path, "chore: add EdgeTest AI generated workflow", yaml_content)
    except GithubException as exc:
        msg = exc.data.get("message", str(exc)) if isinstance(exc.data, dict) else str(exc)
        raise HTTPException(status_code=502, detail=f"GitHub error: {msg}")

    return {"status": "ok", "path": f"{payload.repo_full_name}/{workflow_path}"}


# ---------------------------------------------------------------------------
# GET /export/{session_id}/xlsx
# ---------------------------------------------------------------------------

@router.get("/{session_id}/xlsx")
async def export_xlsx(
    session_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Response:
    session = await db.get(SessionModel, session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    result = await db.execute(
        select(TestRun).where(TestRun.session_id == session_id).order_by(TestRun.created_at.desc()).limit(1)
    )
    test_run = result.scalar_one_or_none()

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise HTTPException(status_code=501, detail="openpyxl not installed — run: pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Unit Tests"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1E3A5F")
    integ_fill = PatternFill("solid", fgColor="3B1F6F")

    # ── Unit Tests sheet ──────────────────────────────────────────────────────
    unit_headers = ["#", "File", "Language", "Lines", "AAA Compliant", "Coverage Warning"]
    for col, h in enumerate(unit_headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    unit_files: list[dict] = session.unit_test_files or []
    for i, f in enumerate(unit_files, 1):
        code = f.get("code", "")
        lines = len(code.splitlines()) if code else 0
        ws.append([
            i,
            f.get("filename", f"unit_test_{i}"),
            f.get("language", session.language),
            lines,
            "Yes" if f.get("aaa_compliant") else "No",
            f.get("classification_warning", ""),
        ])

    # ── Integration Tests sheet ───────────────────────────────────────────────
    ws_integ = wb.create_sheet("Integration Tests")
    integ_headers = ["#", "File", "Language", "Lines", "AAA Compliant", "Coverage Warning"]
    for col, h in enumerate(integ_headers, 1):
        cell = ws_integ.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = integ_fill
        cell.alignment = Alignment(horizontal="center")

    integ_files: list[dict] = session.integration_test_files or []
    for i, f in enumerate(integ_files, 1):
        code = f.get("code", "")
        lines = len(code.splitlines()) if code else 0
        ws_integ.append([
            i,
            f.get("filename", f"integ_test_{i}"),
            f.get("language", session.language),
            lines,
            "Yes" if f.get("aaa_compliant") else "No",
            f.get("classification_warning", ""),
        ])

    # ── Sandbox Results sheet (if available) ─────────────────────────────────
    failures: dict[str, str] = {}
    if test_run and test_run.results_json:
        for fail in test_run.results_json.get("failures", []):
            failures[fail.get("test_name", "")] = fail.get("error_message", "")

    if test_run:
        ws_results = wb.create_sheet("Sandbox Results")
        res_headers = ["#", "Test Name", "Status", "Error Message"]
        for col, h in enumerate(res_headers, 1):
            cell = ws_results.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
        for i in range(1, test_run.passed + 1):
            ws_results.append([i, f"test_{str(i).zfill(3)}", "PASSED", ""])
        for j, (name, msg) in enumerate(failures.items(), test_run.passed + 1):
            ws_results.append([j, name, "FAILED", msg])

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric", "Value"])
    ws2.append(["Session ID", str(session_id)])
    ws2.append(["Language", session.language])
    ws2.append(["Risk Level", session.risk_level or "N/A"])
    ws2.append(["Risk Score", session.risk_score or "N/A"])
    ws2.append(["Unit Test Files", len(unit_files)])
    ws2.append(["Integration Test Files", len(integ_files)])
    ws2.append(["Unit Coverage %", session.unit_coverage_pct or "N/A"])
    ws2.append(["Integration Coverage %", session.integration_coverage_pct or "N/A"])
    if test_run:
        ws2.append(["Total Tests Run", test_run.total_tests])
        ws2.append(["Passed", test_run.passed])
        ws2.append(["Failed", test_run.failed])
        pass_rate = round(test_run.passed / test_run.total_tests * 100, 1) if test_run.total_tests > 0 else 0
        ws2.append(["Pass Rate", f"{pass_rate}%"])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=edgetest-results-{str(session_id)[:8]}.xlsx"},
    )


# ---------------------------------------------------------------------------
# GET /export/{session_id}/json
# ---------------------------------------------------------------------------

@router.get("/{session_id}/json")
async def export_json(
    session_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Response:
    session = await db.get(SessionModel, session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    result = await db.execute(
        select(TestRun).where(TestRun.session_id == session_id).order_by(TestRun.created_at.desc()).limit(1)
    )
    test_run = result.scalar_one_or_none()

    unit_files = session.unit_test_files or []
    integ_files = session.integration_test_files or []

    payload = {
        "session_id": str(session_id),
        "language": session.language,
        "risk_score": session.risk_score,
        "risk_level": session.risk_level,
        "risk_analysis": session.risk_json,
        "traceability_map": session.traceability_map,
        "unit_test_files": [
            {
                "filename": f.get("filename", ""),
                "language": f.get("language", session.language),
                "code": f.get("code", ""),
                "aaa_compliant": f.get("aaa_compliant"),
                "aaa_compliance_percent": f.get("aaa_compliance_percent"),
            }
            for f in unit_files
        ],
        "integration_test_files": [
            {
                "filename": f.get("filename", ""),
                "language": f.get("language", session.language),
                "code": f.get("code", ""),
                "aaa_compliant": f.get("aaa_compliant"),
                "aaa_compliance_percent": f.get("aaa_compliance_percent"),
                "misclassified": f.get("misclassified"),
            }
            for f in integ_files
        ],
        "unit_coverage_pct": session.unit_coverage_pct,
        "integration_coverage_pct": session.integration_coverage_pct,
        "test_results": test_run.results_json if test_run else None,
        "summary": {
            "total": test_run.total_tests if test_run else 0,
            "passed": test_run.passed if test_run else 0,
            "failed": test_run.failed if test_run else 0,
            "unit_files": len(unit_files),
            "integration_files": len(integ_files),
        },
    }

    return Response(
        content=json.dumps(payload, indent=2, default=str),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename=edgetest-results-{str(session_id)[:8]}.json"},
    )


# ---------------------------------------------------------------------------
# GET /export/{session_id}/docx
# ---------------------------------------------------------------------------

@router.get("/{session_id}/docx")
async def export_docx(
    session_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Response:
    session = await db.get(SessionModel, session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found")
    if session.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")

    result = await db.execute(
        select(TestRun).where(TestRun.session_id == session_id).order_by(TestRun.created_at.desc()).limit(1)
    )
    test_run = result.scalar_one_or_none()

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed — run: pip install python-docx")

    doc = Document()
    doc.add_heading("EdgeTest AI — Test Documentation", 0)

    doc.add_heading("Session Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Property"
    hdr[1].text = "Value"

    unit_files = session.unit_test_files or []
    integ_files = session.integration_test_files or []

    rows_data = [
        ("Session ID", str(session_id)),
        ("Language", session.language),
        ("Risk Level", session.risk_level or "N/A"),
        ("Risk Score", str(session.risk_score or "N/A")),
        ("Unit Test Files", str(len(unit_files))),
        ("Integration Test Files", str(len(integ_files))),
        ("Unit Coverage %", str(session.unit_coverage_pct or "N/A")),
        ("Integration Coverage %", str(session.integration_coverage_pct or "N/A")),
    ]
    if test_run:
        pass_rate = round(test_run.passed / test_run.total_tests * 100, 1) if test_run.total_tests > 0 else 0
        rows_data += [
            ("Total Tests Run", str(test_run.total_tests)),
            ("Passed", str(test_run.passed)),
            ("Failed", str(test_run.failed)),
            ("Pass Rate", f"{pass_rate}%"),
        ]
    for k, v in rows_data:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[1].text = v

    if session.risk_json:
        doc.add_heading("Risk Analysis", level=1)
        doc.add_paragraph(session.risk_json.get("human_readable_reason", ""))
        factors = session.risk_json.get("risk_factors", [])
        if factors:
            doc.add_paragraph("Risk Factors:", style="Heading 3")
            for f in factors:
                doc.add_paragraph(f, style="List Bullet")

    if unit_files:
        doc.add_heading("Unit Test Files", level=1)
        for i, f in enumerate(unit_files, 1):
            fname = f.get("filename", f"unit_test_{i}")
            aaa = "AAA Compliant" if f.get("aaa_compliant") else "Not AAA"
            doc.add_paragraph(f"File {i}: {fname}  [{aaa}]", style="Heading 3")
            code = f.get("code", "")
            if code:
                p = doc.add_paragraph(code[:2000] + ("…" if len(code) > 2000 else ""))
                try:
                    p.style = doc.styles["No Spacing"]
                except KeyError:
                    pass

    if integ_files:
        doc.add_heading("Integration Test Files", level=1)
        for i, f in enumerate(integ_files, 1):
            fname = f.get("filename", f"integ_test_{i}")
            warn = " ⚠ misclassified" if f.get("misclassified") else ""
            doc.add_paragraph(f"File {i}: {fname}{warn}", style="Heading 3")
            code = f.get("code", "")
            if code:
                p = doc.add_paragraph(code[:2000] + ("…" if len(code) > 2000 else ""))
                try:
                    p.style = doc.styles["No Spacing"]
                except KeyError:
                    pass

    if test_run and test_run.results_json:
        doc.add_heading("Sandbox Test Results", level=1)
        failures = test_run.results_json.get("failures", [])
        if failures:
            doc.add_heading("Failures", level=2)
            for f in failures:
                doc.add_paragraph(f"✗ {f.get('test_name', '')}: {f.get('error_message', '')}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=edgetest-docs-{str(session_id)[:8]}.docx"},
    )
